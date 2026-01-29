import base64
import os
import re
import uuid
import time
from datetime import datetime
from glob import glob
from typing import Optional, List, Literal, Dict, Tuple

from fastapi import FastAPI, Header, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt

# ======================================================
# CONFIGURACIÓN
# ======================================================
# Token de seguridad (debe coincidir con Render y GPT Action)
API_BEARER_TOKEN = os.getenv("DOCGEN_API_TOKEN", "sk-docgen-change-me")

# Directorio temporal: En Render usamos /tmp porque es el único garantizado para escribir
OUTPUT_DIR = os.getenv("OUTPUT_DIR", "/tmp/docgen_output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# URL pública para construir links absolutos (evita que el GPT invente rutas)
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "").rstrip("/")

# Tiempo de vida de los archivos (24h por defecto)
DOWNLOAD_TTL_SECONDS = int(os.getenv("DOWNLOAD_TTL_SECONDS", "86400"))

app = FastAPI(
    title="Document Generator API",
    version="1.3.0",
    description="Generates DOCX documents with robust persistence (glob fix)"
)

# ======================================================
# MODELOS DE DATOS (Pydantic)
# ======================================================
DocumentType = Literal["cover_letter", "cv", "recruiter_scorecard"]
FormatType = Literal["docx"]

class Candidate(BaseModel):
    full_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None

class Section(BaseModel):
    heading: Optional[str] = None
    text: str

class Content(BaseModel):
    body_markdown: str = Field(..., description="Contenido principal en Markdown o texto")
    sections: Optional[List[Section]] = None

class CreateDocumentRequest(BaseModel):
    document_type: DocumentType
    format: FormatType
    language: str
    title: str
    candidate: Optional[Candidate] = None
    content: Content

class CreateDocumentResponse(BaseModel):
    file_id: str
    file_name: str
    content_type: str
    download_url: str
    download_markdown: str
    file_base64: Optional[str] = None

# ======================================================
# SEGURIDAD
# ======================================================
def verify_bearer_token(authorization: Optional[str]) -> None:
    if not authorization:
        raise HTTPException(status_code=401, detail="Missing Authorization header")
    expected = f"Bearer {API_BEARER_TOKEN}"
    if authorization.strip() != expected:
        raise HTTPException(status_code=401, detail="Unauthorized")

# ======================================================
# UTILIDADES (DOCX + Archivos)
# ======================================================
def sanitize_filename(name: str) -> str:
    name = name.strip()
    name = re.sub(r"[^\w\-\. ]+", "", name, flags=re.UNICODE)
    name = re.sub(r"\s+", "_", name)
    return name[:120]

def add_text_with_basic_markdown(doc: Document, text: str) -> None:
    """Convierte Markdown básico (#, ##, -) a formato Word."""
    if not text: return
    lines = text.replace("\r\n", "\n").split("\n")
    bullet_buffer: List[str] = []

    def flush_bullets():
        nonlocal bullet_buffer
        for bullet in bullet_buffer:
            p = doc.add_paragraph(bullet, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(11)
        bullet_buffer = []

    for line in lines:
        line = line.rstrip()
        if not line.strip():
            flush_bullets()
            doc.add_paragraph("")
            continue
        
        # Títulos
        if line.startswith("## "):
            flush_bullets()
            doc.add_heading(line[3:], level=2)
            continue
        if line.startswith("# "):
            flush_bullets()
            doc.add_heading(line[2:], level=1)
            continue
        
        # Listas
        if line.startswith("- ") or line.startswith("* "):
            bullet_buffer.append(line[2:])
            continue
        
        # Texto normal
        flush_bullets()
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)
    
    flush_bullets()

def cleanup_expired_files() -> None:
    """Borra ficheros antiguos buscando directamente en disco (Robustez ante reinicios)."""
    now = time.time()
    # Busca todos los .docx en la carpeta de salida
    for p in glob(os.path.join(OUTPUT_DIR, "*.docx")):
        try:
            age = now - os.path.getmtime(p)
            if age > DOWNLOAD_TTL_SECONDS:
                os.remove(p)
        except Exception:
            pass

def build_docx(req: CreateDocumentRequest) -> Tuple[str, str, str]:
    doc = Document()
    
    # 1. Título y Metadatos
    doc.add_heading(req.title, level=1)
    meta = []
    if req.candidate and req.candidate.full_name:
        meta.append(req.candidate.full_name)
    meta.append(datetime.now().strftime("%Y-%m-%d"))
    meta.append(f"Language: {req.language}")
    
    p = doc.add_paragraph(" | ".join(meta))
    for run in p.runs:
        run.font.size = Pt(9)
        run.italic = True
    doc.add_paragraph("") # Espacio

    # 2. Contenido Principal
    add_text_with_basic_markdown(doc, req.content.body_markdown)

    # 3. Secciones Extra (Tablas simuladas, preguntas, etc.)
    if req.content.sections:
        doc.add_page_break()
        for section in req.content.sections:
            if section.heading:
                doc.add_heading(section.heading, level=2)
            add_text_with_basic_markdown(doc, section.text)

    # 4. Guardar archivo con ID determinista en el nombre
    safe_title = sanitize_filename(req.title)
    file_id = uuid.uuid4().hex
    # IMPORTANTE: El nombre incluye el ID para poder buscarlo con glob después
    file_name = f"{safe_title}_{file_id}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)
    
    doc.save(file_path)
    return file_id, file_name, file_path

# ======================================================
# ENDPOINTS
# ======================================================

@app.get("/health")
def health():
    return {
        "status": "ok", 
        "version": "1.3.0", 
        "engine": "glob-fix",
        "output_dir": OUTPUT_DIR,
        "public_base_url": PUBLIC_BASE_URL
    }

@app.post("/v1/documents", response_model=CreateDocumentResponse)
def create_document(req: CreateDocumentRequest, authorization: Optional[str] = Header(None)):
    verify_bearer_token(authorization)
    cleanup_expired_files() # Limpieza preventiva
    
    # Generar DOCX
    file_id, file_name, file_path = build_docx(req)
    
    # Construir URL absoluta
    base_url = PUBLIC_BASE_URL if PUBLIC_BASE_URL else ""
    download_path = f"/v1/download/{file_id}"
    download_url = f"{base_url}{download_path}"
    
    # Markdown listo para el GPT (fallback)
    download_markdown = f"[Download DOCX]({download_url})"

    # Base64 opcional (para debugging o descarga directa en tests)
    with open(file_path, "rb") as f:
        encoded = base64.b64encode(f.read()).decode("utf-8")

    return CreateDocumentResponse(
        file_id=file_id,
        file_name=file_name,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        download_url=download_url,
        download_markdown=download_markdown,
        file_base64=encoded
    )

@app.get("/v1/download/{file_id}")
def download_document(file_id: str):
    """
    Busca el archivo en disco usando glob. 
    Esto arregla el error 'unknown id' si el servidor se reinicia.
    """
    cleanup_expired_files()
    
    # Patrón de búsqueda: cualquier nombre que termine en _{file_id}.docx
    # Esto encuentra el archivo aunque no sepamos el título exacto, solo el ID
    pattern = os.path.join(OUTPUT_DIR, f"*_{file_id}.docx")
    matches = glob(pattern)
    
    if not matches:
        raise HTTPException(status_code=404, detail="File not found (expired or invalid ID). Please regenerate.")
    
    # Tomamos el primer archivo que coincida
    file_path = matches
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File lost from disk.")

    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(file_path)
    )
